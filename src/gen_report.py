"""gen_report.py — 產生 report.docx（Phase 1-6 驗收結果與觀察）"""

import datetime
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "report.docx"


# ─────────────────────────────────────────────
# 樣式工具函式
# ─────────────────────────────────────────────

def set_font(run, size=11, bold=False, font_name="新細明體"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font_name
    try:
        run.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        pass


def add_heading(doc, text, level=1):
    h = doc.add_heading(level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.add_run(text)
    sizes = {1: 16, 2: 13, 3: 12}
    set_font(run, size=sizes.get(level, 12), bold=True, font_name="標楷體")
    return h


def add_para(doc, text, indent_cm=0, size=11):
    p = doc.add_paragraph()
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(text)
    set_font(run, size=size)
    return p


def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_font(run, size=size)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # 表頭
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=10, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 資料列
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=10)
    return table


# ─────────────────────────────────────────────
# 主程式
# ─────────────────────────────────────────────

def main():
    doc = Document()

    # 頁面邊距
    sec = doc.sections[0]
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.5)

    # ══ 封面 ══
    doc.add_paragraph()
    doc.add_paragraph()
    for text, size in [
        ("臺北市道路交通事故熱點風險預測", 20),
        ("與照相設備設置建議系統", 20),
        ("", 12),
        ("開發進度報告（Phase 1–6 驗收結果）", 14),
        ("", 12),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_font(run, size=size, bold=(size >= 18), font_name="標楷體")

    date_str = datetime.date.today().strftime("%Y 年 %m 月 %d 日")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("報告日期：" + date_str)
    set_font(run, size=12)
    doc.add_page_break()

    # ══ 一、專案概述 ══
    add_heading(doc, "一、專案概述", 1)
    add_para(doc, (
        "本系統為批次執行之資料分析管線，輸入臺北市政府開放資料（110–114 年交通事故斑點圖、"
        "固定式照相設備一覽表），輸出事故空間熱點清單、時間序列預測模型驗證指標、"
        "照相設備建議設置清單及視覺化圖表。"
    ))
    add_para(doc, (
        "開發語言：Python 3.10+，採模組化架構（s01–s08），所有參數集中於 config.py 管理，"
        "random_state=42 確保可重現性（NFR-1）。"
    ))

    # ══ 二、各 Phase 驗收結果 ══
    add_heading(doc, "二、各 Phase 驗收結果", 1)

    # ── Phase 1 ──
    add_heading(doc, "Phase 1：資料載入與清理（s01_clean.py，FR-01）", 2)
    add_para(doc, (
        "以 cp950 編碼讀取五份事故斑點圖（IN-1~IN-5）並縱向合併，統一欄名為 longitude / latitude"
        "（符合 kepler.gl 格式），依序執行三步驟清理（座標 NaN → 座標超出範圍 → 時間解析失敗）。"
    ))
    add_table(doc,
        ["驗收條件", "預期值", "實際結果", "狀態"],
        [
            ["AC-01a 清理後筆數", "118,980 筆", "118,980 筆", "✓ 通過"],
            ["AC-01b 異常紀錄輸出", "含異常原因之 CSV", "anomalies.csv（3 筆）", "✓ 通過"],
            ["照相設備表筆數", "143 台", "143 台（其中測速 98 台）", "✓ 通過"],
        ]
    )
    doc.add_paragraph()
    add_para(doc, "異常紀錄明細（3 筆，全部為座標欄位 NaN）：")
    add_table(doc,
        ["發生時間", "肇事地點", "異常原因"],
        [
            ["2023/12/13 20:51", "文山區木柵路3段34號", "座標欄位為空值（NaN）"],
            ["2025/3/15 13:22", "大同區太原路50號", "座標欄位為空值（NaN）"],
            ["2025/12/13 14:27", "大安區信義路3段與復興南路口三角公園", "座標欄位為空值（NaN）"],
        ]
    )
    doc.add_paragraph()
    add_para(doc, (
        "【觀察】三筆異常均為座標欄位完全缺漏，無法補值，直接剔除。"
        "清理流程 Step B（座標超出範圍）與 Step C（時間解析失敗）均無觸發，"
        "顯示原始政府資料整體品質良好，僅極少數紀錄缺漏座標。"
    ))

    # ── Phase 2 ──
    add_heading(doc, "Phase 2：空間熱點辨識（s02_hotspot.py，FR-02）", 2)
    add_para(doc, (
        "對事故座標執行 DBSCAN（metric=haversine、eps=50m、min_samples=80、algorithm=ball_tree），"
        "計算各熱點中心座標、事故統計摘要與照相設備距離。"
    ))
    add_table(doc,
        ["驗收條件", "預期值", "實際結果", "狀態"],
        [
            ["AC-02a 熱點數", "145 個", "145 個", "✓ 通過"],
            ["AC-02a 熱點內事故占比", "19.2%（±0.1%）", "19.2%", "✓ 通過"],
            ["AC-02b 噪音點保留", "保留於主檔", "96,122 筆（cluster=-1）", "✓ 通過"],
            ["AC-08 預驗 Top1 地點", "中正區羅斯福路4段與基隆路4段口", "符合", "✓ 通過"],
            ["AC-08 預驗 Top1 件數", "539 件", "539 件", "✓ 通過"],
            ["AC-08 預驗 Top1 最近設備距離", "915 m", "915 m", "✓ 通過"],
        ]
    )
    doc.add_paragraph()
    add_para(doc, (
        "【觀察】145 個熱點中位設備距離 359m（任意設備）、465m（測速設備），"
        "顯示超過半數熱點已在某種照相設備覆蓋範圍 300m 內。"
        "熱點事故占全市 19.2%，意味著約 1/5 的事故集中在極少數空間熱點，"
        "具備高優先介入價值。"
    ))

    # ── Phase 3 ──
    add_heading(doc, "Phase 3：時間序列化（s03_series.py，FR-03）", 2)
    add_para(doc, (
        "產出三份時間序列產物：全市日事故數序列（1,826 天）、"
        "熱點月事故數 panel（145×60）、各熱點時間指紋向量（145×43）。"
    ))
    add_table(doc,
        ["驗收條件", "預期值", "實際結果", "狀態"],
        [
            ["AC-03 日序列長度", "1,826 天", "1,826 天", "✓ 通過"],
            ["AC-03 無缺日", "True", "True（零件日數 = 0）", "✓ 通過"],
            ["AC-03 Panel shape", "(145, 60)", "(145, 60)", "✓ 通過"],
            ["AC-03 無缺月", "True", "60 個月完整", "✓ 通過"],
            ["時間指紋向量 shape", "(145, 43)", "(145, 43)", "✓ 通過"],
            ["向量列加總驗證", "全部 = 1.0", "min = max = 1.0", "✓ 通過"],
        ]
    )
    doc.add_paragraph()
    add_para(doc, (
        "【觀察一】全市日均事故 65.2 件，最高單日 124 件，"
        "整個分析期間（2021–2025）無零件日（最低 7 件），"
        "臺北市每日均有交通事故發生。"
    ))
    add_para(doc, (
        "【觀察二】COVID-19 三級警戒（2021-05-19~07-26，共 69 天）期間事故數明顯下降，"
        "此結構性斷點在 Phase 4 加入 covid_lv3 虛擬變數控制，"
        "避免汙染模型訓練（SRS A-4）。"
    ))

    # ── Phase 4 ──
    add_heading(doc, "Phase 4：特徵萃取（s04_features.py，FR-04）", 2)
    add_para(doc, (
        "從日序列建立機器學習特徵矩陣（18 欄），"
        "所有特徵以 shift 機制確保無未來資訊洩漏（SRS T-4）。"
    ))
    add_table(doc,
        ["驗收條件", "預期值", "實際結果", "狀態"],
        [
            ["AC-04 特徵矩陣 NaN 數", "0", "0", "✓ 通過"],
            ["特徵矩陣 Shape", "—", "(1798, 18)", "✓ 通過"],
            ["暖機期截斷", "前 28 天", "移除 28 筆", "✓ 通過"],
            ["訓練集筆數", "—", "1,433 筆（2021-01-29~2024-12-31）", "✓ 通過"],
            ["驗證集筆數", "—", "365 筆（2025-01-01~2025-12-31）", "✓ 通過"],
            ["T-4 洩漏驗證", "lag 只用過去資訊", "lag_28[2025-01-01] = daily[2024-12-04] ✓", "✓ 通過"],
            ["COVID dummy 天數", "69 天", "69 天", "✓ 通過"],
        ]
    )
    doc.add_paragraph()
    add_para(doc, "特徵群組說明：")
    add_bullet(doc, "Lag 特徵：lag_1 / lag_7 / lag_14 / lag_28（shift(k) 實作，防洩漏）")
    add_bullet(doc, "移動統計：ma_7 / ma_28 / std_7 / std_28（shift(1).rolling(w)，不含當日）")
    add_bullet(doc, "星期 one-hot：weekday_0~6（0=週一，日曆屬性，無洩漏）")
    add_bullet(doc, "月份：month（整數 1–12）")
    add_bullet(doc, "結構控制：covid_lv3（2021-05-19~07-26=1，從 config 常數派生）")

    # ── Phase 5 ──
    add_heading(doc, "Phase 5：PCA 降維與 K-means 熱點型態分群（s05_pca_kmeans.py，FR-05/06）", 2)
    add_para(doc, (
        "對 145×43 時間指紋向量進行 StandardScaler 標準化後 PCA(n=10)，"
        "再以 K-means（k=2~8，random_state=42，n_init=10）選最佳 k，"
        "輸出各群平均時間分布曲線並命名型態，回填 hotspots.csv 的 cluster_type 欄位。"
    ))
    add_table(doc,
        ["驗收條件", "預期值", "實際結果", "狀態"],
        [
            ["AC-05 解釋變異曲線輸出", "✓", "explained_variance.csv 輸出", "✓ 通過"],
            ["AC-05 前3PC累積解釋變異", "—", "18.6%（< 60%，需報告討論）", "⚠ 需討論"],
            ["AC-06 silhouette 分數表輸出", "✓", "kmeans_metrics.csv（k=2~8 完整）", "✓ 通過"],
            ["AC-06 最佳 k", "—", "k=8（silhouette=0.2451）", "✓ 通過"],
            ["AC-06 每熱點皆有群標籤", "145/145", "145/145（null=0）", "✓ 通過"],
            ["cluster_type 回填", "0 空值", "0 空值", "✓ 通過"],
        ]
    )
    doc.add_paragraph()
    add_para(doc, "熱點型態分布（145 個熱點，k=8 群合併為 4 種型態）：")
    add_table(doc,
        ["型態", "熱點數", "占比", "特徵描述", "建議設備類型"],
        [
            ["全日型",    "60", "41.4%", "分布平坦，無顯著尖峰",       "綜合評估（A1 加權）"],
            ["通勤尖峰型","40", "27.6%", "早07-09 + 晚17-19 雙峰",    "闖紅燈/路口科技執法"],
            ["假日型",    "25", "17.2%", "週末占比相對偏高",            "行人安全導向科技執法"],
            ["夜間型",    "20", "13.8%", "深夜22-03點占比高",          "測速照相"],
        ]
    )
    doc.add_paragraph()
    add_para(doc, (
        "【重要發現】臺北市全部 145 個熱點均呈現通勤高峰特徵（commute 值 0.355–0.481），"
        "PCA 前 3 主成分僅解釋 18.6% 變異，顯示熱點間時間分布高度相似。"
        "此為都市密集交通的均質性特徵，「空間位置差異」遠大於「時間模式差異」。"
        "因固定閾值會將所有熱點歸為通勤尖峰型，本系統改採跨群相對閾值（mean ± σ×factor）"
        "動態命名，確保分群結果具差異性與解釋性。"
    ))

    # ── Phase 6 ──
    add_heading(doc, "Phase 6：預測建模與 80/20 驗證（s06_forecast.py，FR-07）", 2)
    add_para(doc, (
        "建立四個預測模型，嚴格時序切分（訓練 2021-01-29~2024-12-31、"
        "驗證 2025 全年 365 天），輸出 MAE/RMSE 驗證指標。"
    ))
    add_table(doc,
        ["模型", "MAE", "RMSE", "vs Naive 改善率", "說明"],
        [
            ["Naive(lag-7)",               "12.433", "16.110", "—（基準）",   "ŷ(t)=y(t-7)，零訓練成本"],
            ["SARIMA(1,1,2)(1,0,1,7)",     "9.154",  "12.051", "+26.4%",      "s=7，36組AIC網格搜尋，20秒"],
            ["XGBoost",                    "9.275",  "11.702", "+25.4%",      "early_stopping，訓練集尾端10%"],
            ["熱點月 XGBoost（全域）",      "1.301",  "1.702",  "不同任務",   "5220訓練樣本，lag_1m~12m"],
        ]
    )
    doc.add_paragraph()
    add_table(doc,
        ["驗收條件", "預期值", "實際結果", "狀態"],
        [
            ["AC-07a 各模型指標輸出", "✓", "metrics.csv 含所有模型 MAE/RMSE", "✓ 通過"],
            ["AC-07b XGBoost 優於 Naive", "改善率 > 0", "MAE 9.275 < 12.433（+25.4%）", "✓ 通過"],
            ["AC-07c 無未來資訊洩漏", "✓", "Naive[2025-01-01]=daily[2024-12-25]=68 ✓", "✓ 通過"],
        ]
    )
    doc.add_paragraph()
    add_para(doc, (
        "【觀察一】SARIMA(MAE=9.154) 略優於 XGBoost(MAE=9.275)，差異僅 0.12 件/日，"
        "兩者改善率相近（~25-26%）。統計模型在強週期性序列中與機器學習模型旗鼓相當。"
    ))
    add_para(doc, (
        "【觀察二】Naive lag-7 本身已提供不錯基礎，顯示臺北市事故具強烈週週期性。"
        "XGBoost 主要貢獻在於捕捉 lag_1/14/28 及移動平均帶來的中短期趨勢修正。"
    ))
    add_para(doc, (
        "【觀察三】全域熱點月 XGBoost MAE=1.301 件/月，在月均 2.6 件的情境下，"
        "預測誤差約 50%，性能受限於單一熱點月均事故數過低（統計雜訊大）。"
    ))

    # ══ 三、驗收總覽 ══
    add_heading(doc, "三、驗收總覽", 1)
    add_table(doc,
        ["Phase", "模組", "關鍵驗收", "結果", "狀態"],
        [
            ["Phase 1", "s01_clean.py",    "118,980 筆；異常紀錄輸出",         "全部通過", "✓"],
            ["Phase 2", "s02_hotspot.py",  "145熱點；19.2%；Top1=539件/915m", "全部通過", "✓"],
            ["Phase 3", "s03_series.py",   "1826天無缺日；145×60；145×43",    "全部通過", "✓"],
            ["Phase 4", "s04_features.py", "NaN=0；T-4無洩漏；COVID=69天",    "全部通過", "✓"],
            ["Phase 5", "s05_pca_kmeans",  "k=8；145有標籤；PCA前3PC=18.6%", "通過（PCA需討論）", "✓⚠"],
            ["Phase 6", "s06_forecast.py", "XGBoost+25.4%；SARIMA+26.4%",    "全部通過", "✓"],
        ]
    )

    # ══ 四、待開發項目 ══
    add_heading(doc, "四、待開發項目", 1)
    add_table(doc,
        ["Phase", "模組", "主要任務"],
        [
            ["Phase 7", "s07_recommend.py", "覆蓋缺口篩選；依風險排序輸出前10名建議清單（AC-08）"],
            ["Phase 8", "s08_visualize.py", "產出 F1–F9 圖表（PNG ≥150dpi）+ F2 HTML 互動地圖"],
        ]
    )

    # ══ 五、報告撰寫注意事項 ══
    add_heading(doc, "五、報告撰寫注意事項（SRS 規範）", 1)
    notes = [
        "【NFR-5 因果宣稱】所有產出文字不得含因果宣稱。照相設備無設置日期，"
        "僅能宣稱「空間相關性」與「風險預測」，不得說「設置照相設備可降低事故數」。",

        "【AC-05 PCA 討論】前 3 主成分僅解釋 18.6%（< 60%），需在報告說明原因："
        "臺北市熱點時間分布高度相似（均為通勤高峰），43 維向量無明顯線性主軸，"
        "因此 PCA 效果有限，但 K-means 仍能從高維空間找出有意義的群組。",

        "【資料不平衡】A1 事故僅約 0.3%，本系統以 A1+A2 合計為預測目標，"
        "A1 作排序加權，報告需說明此設計理由（避免過稀疏目標變數的模型不穩定）。",

        "【排序限制】未以車流量正規化（SRS 附錄限制 4），建議清單排序屬「絕對風險」"
        "而非「相對風險率」，應在報告限制章節說明。",

        "【資料來源標注】報告需載明各資料集名稱與下載位置："
        "臺北市資料大平臺（data.taipei）/ 政府資料開放平臺（data.gov.tw）/ "
        "中央氣象署 CODiS（codis.cwa.gov.tw）。",
    ]
    for note in notes:
        add_bullet(doc, note)

    doc.save(str(OUTPUT))
    print("report.docx 產生完成：", OUTPUT)


if __name__ == "__main__":
    main()
